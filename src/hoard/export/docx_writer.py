"""docx_writer.py — DOCX export via python-docx.

Takes the compiled Markdown report body and produces a properly formatted
Word document with:
    - Cover page (project name, date, jurisdiction)
    - Heading-styled sections (Heading 1 / 2 / 3)
    - Body paragraphs with justified text, 11 pt Calibri
    - Inline images at their Markdown references
    - Appendix tables from the context register / finds / sample registers

Replaces the earlier pandoc-based DOCX path which required a system-wide
pandoc installation.

exports: write_docx(md_text, path, config, template_yaml) -> Path
used_by: hoard.phases.phase5  → export_report()
rules:   No GPU, no subprocess calls. Pure python-docx.
license: MIT
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore[import-untyped]
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import-untyped]


def write_docx(
    md_text: str,
    output_path: Path,
    project_name: str = "Excavation Report",
    jurisdiction: str = "",
    appendices: dict[str, str] | None = None,
    template_yaml: dict[str, Any] | None = None,
) -> Path:
    """Write the compiled Markdown report to a formatted Word document.

    Args:
        md_text: Full Markdown report text.
        output_path: Where to save the .docx.
        project_name: Project name for the cover page.
        jurisdiction: Jurisdiction label.
        appendices: Optional dict of {appendix_id: markdown_table_text}.
        template_yaml: Optional jurisdiction template config for style overrides.

    Returns:
        output_path (confirmed written).
    """
    doc = Document()

    _set_default_styles(doc)

    # ── Cover page ──────────────────────────────────────────────────────────
    _add_cover_page(doc, project_name, jurisdiction)

    # ── Parse Markdown body ─────────────────────────────────────────────────
    _add_markdown_body(doc, md_text)

    # ── Appendices ──────────────────────────────────────────────────────────
    if appendices:
        _add_appendices(doc, appendices)

    # ── Save ────────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# Internal helpers
# ═══════════════════════════════════════════════════════════════════════════════


def _set_default_styles(doc: Document) -> None:
    """Configure sensible default styles for an archaeological report."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in (1, 2, 3):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        if level == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        else:
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)


def _add_cover_page(doc: Document, project_name: str, jurisdiction: str) -> None:
    """Add a simple cover page: title, subtitle, date."""
    for _ in range(6):
        doc.add_paragraph("")  # Vertical space

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(project_name)
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    if jurisdiction:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(f"Prepared in accordance with {jurisdiction}")
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(24)
    date_run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %Y"))
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def _add_markdown_body(doc: Document, md_text: str) -> None:
    """Parse Markdown and append formatted sections to the document.

    Handles:
        - # / ## / ### headings
        - Body paragraphs (blank-line delimited)
        - Inline images ![alt](path)
        - Markdown tables
        - Bold (**text**) and italic (*text*)
        - --- horizontal rules as page breaks
        - Bullet lists (- / *)
    """
    lines = md_text.split("\n")
    i = 0
    table_rows: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Horizontal rule → page break ──
        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        # ── Heading ──
        if stripped.startswith("#"):
            _flush_table(doc, table_rows)
            table_rows = []
            heading_level = len(stripped.split(" ")[0])  # count #
            if heading_level > 3:
                heading_level = 3
            text = stripped.lstrip("# ").strip()
            doc.add_heading(text, level=heading_level)
            i += 1
            continue

        # ── Bullet list ──
        if stripped.startswith("- ") or stripped.startswith("* "):
            _flush_table(doc, table_rows)
            table_rows = []
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_run_with_formatting(p, text)
            i += 1
            continue

        # ── Table ──
        if "|" in stripped and stripped.startswith("|"):
            table_rows.append(stripped)
            i += 1
            continue

        # ── Image ──
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            _flush_table(doc, table_rows)
            table_rows = []
            alt_text = img_match.group(1)
            img_path_str = img_match.group(2)
            img_path = Path(img_path_str)
            if img_path.exists():
                max_width = doc.sections[0].page_width.inches - doc.sections[0].left_margin.inches - doc.sections[0].right_margin.inches
                doc.add_picture(str(img_path), width=Inches(min(5.5, max_width)))
                # Add caption
                cap = doc.add_paragraph(alt_text)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.italic = True
                cap.runs[0].font.size = Pt(9)
            i += 1
            continue

        # ── YAML front-matter skip ──
        if stripped in ("---", "...") and i < 5:
            # Skip front matter block
            if stripped == "---":
                i += 1
                while i < len(lines) and lines[i].strip() != "---":
                    i += 1
                i += 1  # skip close
                continue

        # ── Empty line → paragraph break ──
        if not stripped:
            _flush_table(doc, table_rows)
            table_rows = []
            i += 1
            continue

        # ── Regular paragraph ──
        _flush_table(doc, table_rows)
        table_rows = []
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_run_with_formatting(p, stripped)
        i += 1

    # Flush any remaining table
    _flush_table(doc, table_rows)


def _add_run_with_formatting(p: Any, text: str) -> None:
    """Add runs to a paragraph, handling **bold** and *italic* markers."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("***"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def _flush_table(doc: Document, rows: list[str]) -> None:
    """Convert accumulated Markdown table rows into a Word table."""
    if not rows:
        return

    # Filter separator row (e.g. |---|---|)
    data_rows: list[list[str]] = []
    for row in rows:
        cells = [c.strip() for c in row.split("|")[1:-1]]
        # Skip separator rows
        if all(re.match(r"^-+$", c) for c in cells if c):
            continue
        data_rows.append(cells)

    if not data_rows:
        return

    table = doc.add_table(rows=len(data_rows), cols=len(data_rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            # Bold the header row
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph("")  # spacing after table


def _add_appendices(doc: Document, appendices: dict[str, str]) -> None:
    """Add appendix sections from the provided markdown texts."""
    doc.add_page_break()
    doc.add_heading("Appendices", level=1)

    labels = {
        "context_register": "Appendix 1: Context Register",
        "finds_concordance": "Appendix 2: Finds Concordance",
        "sample_register": "Appendix 3: Sample Register",
    }

    for app_id, label in labels.items():
        content = appendices.get(app_id, "")
        if not content or content.startswith("*No"):
            continue

        doc.add_heading(label, level=2)
        _add_markdown_body(doc, content)
